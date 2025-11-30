from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import os
import sys
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import io
import json
from docx import Document
from docx.shared import Pt

# When packaged with PyInstaller, resources are extracted to sys._MEIPASS.
# Use that as base; otherwise use the project directory.
BASE_DIR = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

# Create Flask with explicit template/static folders so bundled exe can find them
app = Flask(__name__, template_folder=os.path.join(BASE_DIR, 'app', 'templates'), static_folder=os.path.join(BASE_DIR, 'app', 'static'))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(BASE_DIR, 'troskovi.db')
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'app', 'uploads')
# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

db = SQLAlchemy(app)

# Ensure app data directory exists for per-user storage
APP_DATA_DIR = os.path.join(BASE_DIR, 'app', 'data')
os.makedirs(APP_DATA_DIR, exist_ok=True)

def ensure_user_folder(osoba_id: int):
    user_dir = os.path.join(APP_DATA_DIR, f'user_{osoba_id}')
    hist_dir = os.path.join(user_dir, 'history')
    os.makedirs(hist_dir, exist_ok=True)
    return user_dir, hist_dir

def save_month_snapshot(osoba_id: int, mesec: str):
    """Save a JSON snapshot of incomes and expenses for the given user and month."""
    prihodi_list = Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    troskovi_list = Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()

    prihodi = [{'id': p.id, 'naziv': p.naziv, 'iznos': p.iznos} for p in prihodi_list]
    troskovi = [{'id': t.id, 'naziv': t.naziv, 'iznos': t.iznos, 'kategorija': t.kategorija.naziv} for t in troskovi_list]

    user_dir, hist_dir = ensure_user_folder(osoba_id)
    filename = os.path.join(hist_dir, f'{mesec}.json')
    snapshot = {
        'osoba_id': osoba_id,
        'mesec': mesec,
        'prihodi': prihodi,
        'troskovi': troskovi,
        'generated_at': datetime.utcnow().isoformat()
    }
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(snapshot, f, ensure_ascii=False, indent=2)
    except Exception:
        # don't break main flow if write fails
        pass

# Database Models
class Person(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ime = db.Column(db.String(100), nullable=False)
    prezime = db.Column(db.String(100), nullable=False)
    datum_rodjenja = db.Column(db.Date, nullable=False)
    slika = db.Column(db.String(255))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    prihodi = db.relationship('Prihod', backref='osoba', lazy=True, cascade='all, delete-orphan')
    troskovi = db.relationship('Trosak', backref='osoba', lazy=True, cascade='all, delete-orphan')
    meseci = db.relationship('MesecniIzvestaj', backref='osoba', lazy=True, cascade='all, delete-orphan')

class Prihod(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    osoba_id = db.Column(db.Integer, db.ForeignKey('person.id'), nullable=False)
    naziv = db.Column(db.String(100), nullable=False)
    iznos = db.Column(db.Float, nullable=False)
    mesec = db.Column(db.String(7))  # YYYY-MM format
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class TrosakKategorija(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    naziv = db.Column(db.String(100), nullable=False, unique=True)
    boja = db.Column(db.String(7), default='#3498db')
    
    troskovi = db.relationship('Trosak', backref='kategorija', lazy=True, cascade='all, delete-orphan')

class Trosak(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    osoba_id = db.Column(db.Integer, db.ForeignKey('person.id'), nullable=False)
    kategorija_id = db.Column(db.Integer, db.ForeignKey('trosak_kategorija.id'), nullable=False)
    naziv = db.Column(db.String(100), nullable=False)
    iznos = db.Column(db.Float, nullable=False)
    mesec = db.Column(db.String(7))  # YYYY-MM format
    opis = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class MesecniIzvestaj(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    osoba_id = db.Column(db.Integer, db.ForeignKey('person.id'), nullable=False)
    mesec = db.Column(db.String(7), nullable=False)  # YYYY-MM format
    ukupno_prihodi = db.Column(db.Float, default=0)
    ukupno_troskovi = db.Column(db.Float, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def calculate_percentages(troskovi):
    """Kalkula procentualne vrednosti troškova"""
    if not troskovi:
        return {}
    
    total = sum(t['iznos'] for t in troskovi)
    if total == 0:
        return {}
    
    result = {}
    for trosak in troskovi:
        result[trosak['id']] = (trosak['iznos'] / total) * 100
    return result

# Routes
@app.route('/')
def index():
    osoba = Person.query.first()
    return render_template('index.html', osoba=osoba)

@app.route('/api/osoba', methods=['GET', 'POST'])
def osoba():
    if request.method == 'POST':
        data = request.form
        slika_filename = None
        
        # Handle image upload
        if 'slika' in request.files:
            file = request.files['slika']
            if file and file.filename and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
                filename = timestamp + filename
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                slika_filename = filename
        
        # Update or create person
        osoba = Person.query.first()
        if osoba:
            osoba.ime = data.get('ime')
            osoba.prezime = data.get('prezime')
            osoba.datum_rodjenja = datetime.strptime(data.get('datum_rodjenja'), '%Y-%m-%d').date()
            if slika_filename:
                osoba.slika = slika_filename
        else:
            osoba = Person(
                ime=data.get('ime'),
                prezime=data.get('prezime'),
                datum_rodjenja=datetime.strptime(data.get('datum_rodjenja'), '%Y-%m-%d').date(),
                slika=slika_filename
            )
            db.session.add(osoba)
        
        db.session.commit()
        return jsonify({'success': True, 'osoba_id': osoba.id})
    
    osoba = Person.query.first()
    if osoba:
        return jsonify({
            'id': osoba.id,
            'ime': osoba.ime,
            'prezime': osoba.prezime,
            'datum_rodjenja': osoba.datum_rodjenja.isoformat(),
            'slika': osoba.slika
        })
    return jsonify({}), 404

@app.route('/api/prihodi/<int:osoba_id>/<mesec>', methods=['GET', 'POST'])
def prihodi(osoba_id, mesec):
    if request.method == 'POST':
        data = request.json
        prihod = Prihod(
            osoba_id=osoba_id,
            naziv=data.get('naziv'),
            iznos=float(data.get('iznos')),
            mesec=mesec
        )
        db.session.add(prihod)
        db.session.commit()
        
        # Update izvestaj
        ukupno = sum(p.iznos for p in Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all())
        izvestaj = MesecniIzvestaj.query.filter_by(osoba_id=osoba_id, mesec=mesec).first()
        if izvestaj:
            izvestaj.ukupno_prihodi = ukupno
        db.session.commit()
        
        # Save history snapshot after change
        try:
            save_month_snapshot(osoba_id, mesec)
        except Exception:
            pass

        return jsonify({'success': True, 'prihod_id': prihod.id})
    
    prihodi_list = Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    return jsonify([{
        'id': p.id,
        'naziv': p.naziv,
        'iznos': p.iznos
    } for p in prihodi_list])

@app.route('/api/prihod/<int:prihod_id>', methods=['DELETE'])
def delete_prihod(prihod_id):
    prihod = Prihod.query.get(prihod_id)
    if prihod:
        osoba_id = prihod.osoba_id
        mesec = prihod.mesec
        db.session.delete(prihod)
        db.session.commit()
        
        # Update izvestaj
        ukupno = sum(p.iznos for p in Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all())
        izvestaj = MesecniIzvestaj.query.filter_by(osoba_id=osoba_id, mesec=mesec).first()
        if izvestaj:
            izvestaj.ukupno_prihodi = ukupno
        db.session.commit()
        
        # Save history snapshot after change
        try:
            save_month_snapshot(osoba_id, mesec)
        except Exception:
            pass

        return jsonify({'success': True})
    return jsonify({'success': False}), 404

@app.route('/api/kategorije', methods=['GET', 'POST'])
def kategorije():
    if request.method == 'POST':
        data = request.json
        # Check if category already exists
        existing = TrosakKategorija.query.filter_by(naziv=data.get('naziv')).first()
        if existing:
            return jsonify({'success': False, 'message': 'Kategorija već postoji'}), 400
        
        kategorija = TrosakKategorija(
            naziv=data.get('naziv'),
            boja=data.get('boja', '#3498db')
        )
        db.session.add(kategorija)
        db.session.commit()
        return jsonify({'success': True, 'kategorija_id': kategorija.id})
    
    kategorije_list = TrosakKategorija.query.all()
    return jsonify([{
        'id': k.id,
        'naziv': k.naziv,
        'boja': k.boja
    } for k in kategorije_list])

@app.route('/api/troskovi/<int:osoba_id>/<mesec>', methods=['GET', 'POST'])
def troskovi(osoba_id, mesec):
    if request.method == 'POST':
        data = request.json
        trosak = Trosak(
            osoba_id=osoba_id,
            kategorija_id=int(data.get('kategorija_id')),
            naziv=data.get('naziv'),
            iznos=float(data.get('iznos')),
            mesec=mesec,
            opis=data.get('opis', '')
        )
        db.session.add(trosak)
        db.session.commit()
        
        # Update izvestaj
        ukupno = sum(t.iznos for t in Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all())
        izvestaj = MesecniIzvestaj.query.filter_by(osoba_id=osoba_id, mesec=mesec).first()
        if not izvestaj:
            izvestaj = MesecniIzvestaj(osoba_id=osoba_id, mesec=mesec)
            db.session.add(izvestaj)
        izvestaj.ukupno_troskovi = ukupno
        db.session.commit()
        
        # Save history snapshot after change
        try:
            save_month_snapshot(osoba_id, mesec)
        except Exception:
            pass

        return jsonify({'success': True, 'trosak_id': trosak.id})
    
    troskovi_list = Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    return jsonify([{
        'id': t.id,
        'naziv': t.naziv,
        'iznos': t.iznos,
        'kategorija': t.kategorija.naziv,
        'kategorija_id': t.kategorija_id,
        'kategorija_boja': t.kategorija.boja,
        'opis': t.opis
    } for t in troskovi_list])

@app.route('/api/trosak/<int:trosak_id>', methods=['DELETE'])
def delete_trosak(trosak_id):
    trosak = Trosak.query.get(trosak_id)
    if trosak:
        osoba_id = trosak.osoba_id
        mesec = trosak.mesec
        db.session.delete(trosak)
        db.session.commit()
        
        # Update izvestaj
        ukupno = sum(t.iznos for t in Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all())
        izvestaj = MesecniIzvestaj.query.filter_by(osoba_id=osoba_id, mesec=mesec).first()
        if izvestaj:
            izvestaj.ukupno_troskovi = ukupno
        db.session.commit()
        
        # Save history snapshot after change
        try:
            save_month_snapshot(osoba_id, mesec)
        except Exception:
            pass

        return jsonify({'success': True})
    return jsonify({'success': False}), 404

@app.route('/api/izvestaj/<int:osoba_id>/<mesec>')
def get_izvestaj(osoba_id, mesec):
    osoba = Person.query.get(osoba_id)
    prihodi_list = Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    troskovi_list = Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    
    total_prihodi = sum(p.iznos for p in prihodi_list)
    total_troskovi = sum(t.iznos for t in troskovi_list)
    
    # Group troskovi by kategorija
    troskovi_by_kategorija = {}
    for trosak in troskovi_list:
        kat_naziv = trosak.kategorija.naziv
        if kat_naziv not in troskovi_by_kategorija:
            troskovi_by_kategorija[kat_naziv] = {
                'iznos': 0,
                'boja': trosak.kategorija.boja,
                'procentualno': 0
            }
        troskovi_by_kategorija[kat_naziv]['iznos'] += trosak.iznos
    
    # Calculate percentages
    if total_troskovi > 0:
        for kat in troskovi_by_kategorija:
            troskovi_by_kategorija[kat]['procentualno'] = (troskovi_by_kategorija[kat]['iznos'] / total_troskovi) * 100
    
    return jsonify({
        'osoba': {
            'ime': osoba.ime,
            'prezime': osoba.prezime,
            'slika': osoba.slika
        },
        'mesec': mesec,
        'total_prihodi': total_prihodi,
        'total_troskovi': total_troskovi,
        'razlika': total_prihodi - total_troskovi,
        'troskovi_by_kategorija': troskovi_by_kategorija,
        'broj_stavki_troskova': len(troskovi_list)
    })

@app.route('/api/export-docx/<int:osoba_id>/<mesec>')
def export_docx(osoba_id, mesec):
    osoba = Person.query.get(osoba_id)
    prihodi_list = Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    troskovi_list = Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()

    total_prihodi = sum(p.iznos for p in prihodi_list)
    total_troskovi = sum(t.iznos for t in troskovi_list)

    docx_doc = Document()
    docx_doc.add_heading(f'Mesečni Izveštaj Troškova - {mesec}', level=1)
    docx_doc.add_paragraph(f'Osoba: {osoba.ime} {osoba.prezime}')
    docx_doc.add_paragraph(f'Datum izveštaja: {datetime.now().strftime("%d.%m.%Y")}')
    docx_doc.add_paragraph('')

    # Summary table
    summary = docx_doc.add_table(rows=1, cols=2)
    hdr_cells = summary.rows[0].cells
    hdr_cells[0].text = 'Stavka'
    hdr_cells[1].text = 'Iznos (дин)'
    for label, amount in [('Ukupni Prihodi', total_prihodi), ('Ukupni Troškovi', total_troskovi), ('Razlika (Štednja/Deficit)', total_prihodi - total_troskovi)]:
        row_cells = summary.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = f"{amount:,.2f}"

    docx_doc.add_paragraph('')

    # Prihodi
    if prihodi_list:
        docx_doc.add_heading('PRIHODI', level=2)
        t = docx_doc.add_table(rows=1, cols=2)
        th = t.rows[0].cells
        th[0].text = 'Naziv'
        th[1].text = 'Iznos (дин)'
        for p in prihodi_list:
            r = t.add_row().cells
            r[0].text = p.naziv
            r[1].text = f"{p.iznos:,.2f}"
        r = t.add_row().cells
        r[0].text = 'UKUPNO'
        r[1].text = f"{total_prihodi:,.2f}"

    # Troskovi
    if troskovi_list:
        docx_doc.add_heading('TROŠKOVI', level=2)
        t2 = docx_doc.add_table(rows=1, cols=4)
        th2 = t2.rows[0].cells
        th2[0].text = 'Kategorija'
        th2[1].text = 'Naziv'
        th2[2].text = 'Iznos (дин)'
        th2[3].text = '%'
        for t_item in troskovi_list:
            perc = (t_item.iznos / total_troskovi * 100) if total_troskovi > 0 else 0
            r = t2.add_row().cells
            r[0].text = t_item.kategorija.naziv
            r[1].text = t_item.naziv
            r[2].text = f"{t_item.iznos:,.2f}"
            r[3].text = f"{perc:.1f}%"
        r = t2.add_row().cells
        r[0].text = ''
        r[1].text = 'UKUPNO'
        r[2].text = f"{total_troskovi:,.2f}"
        r[3].text = '100%'

    # Save to user folder
    user_dir, hist_dir = ensure_user_folder(osoba_id)
    filename = os.path.join(user_dir, f'Izvestaj_{osoba.ime}_{osoba.prezime}_{mesec}.docx')
    try:
        docx_doc.save(filename)
    except Exception:
        pass

    return send_file(filename, as_attachment=True)

@app.route('/api/export-pdf/<int:osoba_id>/<mesec>')
def export_pdf(osoba_id, mesec):
    osoba = Person.query.get(osoba_id)
    prihodi_list = Prihod.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    troskovi_list = Trosak.query.filter_by(osoba_id=osoba_id, mesec=mesec).all()
    
    total_prihodi = sum(p.iznos for p in prihodi_list)
    total_troskovi = sum(t.iznos for t in troskovi_list)
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=30,
        alignment=1  # center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#34495e'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title and person info
    elements.append(Paragraph(f"Mesečni Izveštaj Troškova", title_style))
    elements.append(Spacer(1, 0.3 * inch))
    
    # Person info
    person_info = f"<b>Osoba:</b> {osoba.ime} {osoba.prezime}<br/><b>Mesec:</b> {mesec}<br/><b>Datum izveštaja:</b> {datetime.now().strftime('%d.%m.%Y')}"
    elements.append(Paragraph(person_info, styles['Normal']))
    elements.append(Spacer(1, 0.3 * inch))
    
    # Summary table
    elements.append(Paragraph("PREGLED", heading_style))
    summary_data = [
        ['Stavka', 'Iznos (дин)'],
        ['Ukupni Prihodi', f'{total_prihodi:,.2f}'],
        ['Ukupni Troškovi', f'{total_troskovi:,.2f}'],
        ['Razlika (Štednja/Deficit)', f'{total_prihodi - total_troskovi:,.2f}']
    ]
    
    summary_table = Table(summary_data, colWidths=[3*inch, 1.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 0.3 * inch))
    
    # Prihodi table
    if prihodi_list:
        elements.append(Paragraph("PRIHODI", heading_style))
        prihodi_data = [['Naziv', 'Iznos (дин)']]
        for p in prihodi_list:
            prihodi_data.append([p.naziv, f'{p.iznos:,.2f}'])
        prihodi_data.append(['UKUPNO', f'{total_prihodi:,.2f}'])
        
        prihodi_table = Table(prihodi_data, colWidths=[3*inch, 1.5*inch])
        prihodi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#27ae60')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#d5f4e6')),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        elements.append(prihodi_table)
        elements.append(Spacer(1, 0.3 * inch))
    
    # Troškovi table
    if troskovi_list:
        elements.append(Paragraph("TROŠKOVI", heading_style))
        troskovi_data = [['Kategorija', 'Naziv', 'Iznos (дин)', '%']]
        
        total_by_cat = {}
        for t in troskovi_list:
            cat = t.kategorija.naziv
            if cat not in total_by_cat:
                total_by_cat[cat] = 0
            total_by_cat[cat] += t.iznos
        
        for t in troskovi_list:
            percentage = (t.iznos / total_troskovi * 100) if total_troskovi > 0 else 0
            troskovi_data.append([
                t.kategorija.naziv,
                t.naziv,
                f'{t.iznos:,.2f}',
                f'{percentage:.1f}%'
            ])
        
        troskovi_data.append(['', 'UKUPNO', f'{total_troskovi:,.2f}', '100%'])
        
        troskovi_table = Table(troskovi_data, colWidths=[1.2*inch, 1.8*inch, 1.2*inch, 0.8*inch])
        troskovi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e74c3c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#fadbd8')),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        elements.append(troskovi_table)
    
    doc.build(elements)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'Izvestaj_{osoba.ime}_{osoba.prezime}_{mesec}.pdf'
    )

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, host='localhost', port=5000)
